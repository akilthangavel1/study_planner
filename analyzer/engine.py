import re
import io
from datetime import datetime, timedelta
from typing import Dict, List, Any, Optional

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

# Comprehensive Skills Dictionary mapping normalized skill name to aliases/keywords
SKILL_PATTERNS = {
    # Programming Languages
    "Python": [r"\bpython\b", r"\bpython3\b", r"\bpython2\b"],
    "JavaScript": [r"\bjavascript\b", r"\bjs\b", r"\bes6\b", r"\becmascript\b"],
    "TypeScript": [r"\btypescript\b", r"\bts\b"],
    "Java": [r"\bjava\b(?!script)"],
    "C++": [r"\bc\+\+\b", r"\bcpp\b"],
    "C#": [r"\bc#\b", r"\bcsharp\b", r"\bdotnet\b", r"\b\.net\b"],
    "Go": [r"\bgolang\b", r"\bgo\s+language\b", r"\bgo\b"],
    "Rust": [r"\brust\b", r"\brustlang\b"],
    "PHP": [r"\bphp\b"],
    "Ruby": [r"\bruby\b", r"\bruby\s+on\s+rails\b", r"\brails\b"],
    "SQL": [r"\bsql\b", r"\btransact-sql\b", r"\bpl/sql\b"],
    "R": [r"\br\s+programming\b", r"\br\s+language\b", r"\br-project\b"],
    "Swift": [r"\bswift\b", r"\bswiftui\b"],
    "Kotlin": [r"\bkotlin\b"],
    "HTML/CSS": [r"\bhtml\b", r"\bhtml5\b", r"\bcss\b", r"\bcss3\b", r"\bsass\b", r"\bscss\b"],

    # Web & Full-Stack Frameworks
    "Django": [r"\bdjango\b", r"\bdjango\s+rest\s+framework\b", r"\bdrf\b"],
    "FastAPI": [r"\bfastapi\b"],
    "Flask": [r"\bflask\b"],
    "React": [r"\breact\b", r"\breact\.js\b", r"\breactjs\b", r"\breact\s+native\b"],
    "Next.js": [r"\bnext\.js\b", r"\bnextjs\b"],
    "Vue.js": [r"\bvue\b", r"\bvue\.js\b", r"\bvuejs\b", r"\bnuxt\b"],
    "Angular": [r"\bangular\b", r"\bangularjs\b"],
    "Node.js": [r"\bnode\b", r"\bnode\.js\b", r"\bnodejs\b", r"\bexpress\b", r"\bexpress\.js\b"],
    "Spring Boot": [r"\bspring\s*boot\b", r"\bspring\s*framework\b", r"\bspring\b"],
    "Tailwind CSS": [r"\btailwind\b", r"\btailwindcss\b"],
    "Bootstrap": [r"\bbootstrap\b"],
    "GraphQL": [r"\bgraphql\b", r"\bapollo\b"],
    "REST APIs": [r"\brest\b", r"\brestful\b", r"\brest\s+apis?\b", r"\bapi\s+design\b"],

    # Databases & Storage
    "PostgreSQL": [r"\bpostgres\b", r"\bpostgresql\b"],
    "MySQL": [r"\bmysql\b", r"\bmariadb\b"],
    "MongoDB": [r"\bmongodb\b", r"\bmongo\b", r"\bdocumentdb\b"],
    "Redis": [r"\bredis\b", r"\bkey-value\b"],
    "SQLite": [r"\bsqlite\b", r"\bsqlite3\b"],
    "Elasticsearch": [r"\belasticsearch\b", r"\belk\b"],
    "Cassandra": [r"\bcassandra\b"],
    "Firebase": [r"\bfirebase\b", r"\bfirestore\b"],

    # Cloud, DevOps & Infrastructure
    "AWS": [r"\baws\b", r"\bamazon\s+web\s+services\b", r"\bec2\b", r"\bs3\b", r"\blambda\b", r"\brds\b"],
    "Azure": [r"\bazure\b", r"\bmicrosoft\s+azure\b"],
    "Google Cloud (GCP)": [r"\bgcp\b", r"\bgoogle\s+cloud\b"],
    "Docker": [r"\bdocker\b", r"\bcontainerization\b", r"\bcontainers\b"],
    "Kubernetes": [r"\bkubernetes\b", r"\bk8s\b"],
    "Terraform": [r"\bterraform\b", r"\biac\b", r"\binfrastructure\s+as\s+code\b"],
    "CI/CD": [r"\bci/cd\b", r"\bci-cd\b", r"\bgithub\s+actions\b", r"\bjenkins\b", r"\bgitlab\s+ci\b"],
    "Linux / Shell": [r"\blinux\b", r"\bbash\b", r"\bzsh\b", r"\bshell\s+scripting\b", r"\bunix\b"],
    "Git & GitHub": [r"\bgit\b", r"\bgithub\b", r"\bgitlab\b", r"\bbitbucket\b"],
    "Prometheus / Grafana": [r"\bprometheus\b", r"\bgrafana\b", r"\bmonitoring\b", r"\bobservability\b"],

    # AI, Machine Learning & Data Science
    "Machine Learning": [r"\bmachine\s+learning\b", r"\bml\b", r"\bscikit-learn\b", r"\bsklearn\b"],
    "Deep Learning": [r"\bdeep\s+learning\b", r"\bneural\s+networks\b", r"\bcnn\b", r"\brnn\b", r"\blstm\b"],
    "PyTorch": [r"\bpytorch\b", r"\btorch\b"],
    "TensorFlow": [r"\btensorflow\b", r"\bkeras\b", r"\btf\b"],
    "Natural Language Processing (NLP)": [r"\bnlp\b", r"\bnatural\s+language\s+processing\b", r"\bhugging\s*face\b", r"\btransformers\b", r"\bspacy\b", r"\bnltk\b"],
    "Computer Vision (CV)": [r"\bcomputer\s+vision\b", r"\bopencv\b", r"\byolo\b", r"\bimage\s+processing\b"],
    "Generative AI & LLMs": [r"\bllms?\b", r"\blarge\s+language\s+models?\b", r"\bgenai\b", r"\bgenerative\s+ai\b", r"\blangchain\b", r"\bllama-index\b", r"\bprompt\s+engineering\b", r"\brag\b"],
    "Pandas": [r"\bpandas\b"],
    "NumPy": [r"\bnumpy\b"],
    "Data Analysis / Visualization": [r"\bdata\s+analysis\b", r"\bdata\s+analytics\b", r"\bmatplotlib\b", r"\bseaborn\b", r"\bpower\s*bi\b", r"\btableau\b"],

    # Data Engineering & Big Data
    "Apache Spark": [r"\bspark\b", r"\bpyspark\b", r"\bapache\s+spark\b"],
    "Apache Kafka": [r"\bkafka\b", r"\bapache\s+kafka\b", r"\bevent\s+streaming\b"],
    "Data Warehousing": [r"\bsnowflake\b", r"\bbigquery\b", r"\bredshift\b", r"\bdbt\b"],
    "Airflow": [r"\bairflow\b", r"\bapache\s+airflow\b", r"\bdata\s+pipelines\b", r"\betl\b"],

    # Cybersecurity
    "Network Security": [r"\bnetwork\s+security\b", r"\bwireshark\b", r"\btcp/ip\b", r"\bfirewall\b", r"\bvpn\b"],
    "Penetration Testing": [r"\bpenetration\s+testing\b", r"\bpen\s+testing\b", r"\bethical\s+hacking\b", r"\bmetasploit\b", r"\bburp\s+suite\b", r"\bnmap\b"],
    "SIEM & SOC Analysis": [r"\bsiem\b", r"\bsoc\b", r"\bsplunk\b", r"\bqradar\b", r"\bincident\s+response\b"],
    "Cryptography & Identity": [r"\bcryptography\b", r"\boauth\b", r"\bjwt\b", r"\biam\b", r"\bpki\b"],

    # Mobile Development
    "Flutter": [r"\bflutter\b", r"\bdart\b"],
    "React Native": [r"\breact\s+native\b"],
    "iOS Development": [r"\bios\b", r"\bxcode\b", r"\bswift\b", r"\bobjective-c\b"],
    "Android Development": [r"\bandroid\b", r"\bandroid\s+studio\b", r"\bkotlin\b"],

    # UI/UX & Design
    "UI/UX Design": [r"\bui/ux\b", r"\bui\s+design\b", r"\bux\s+design\b", r"\buser\s+experience\b", r"\buser\s+interface\b"],
    "Figma": [r"\bfigma\b", r"\badobe\s+xd\b", r"\bwireframing\b", r"\bprototyping\b"],

    # Core Computer Science & Methodologies
    "Data Structures & Algorithms": [r"\bdsa\b", r"\bdata\s+structures\b", r"\balgorithms\b", r"\bleetcode\b"],
    "Agile / Scrum": [r"\bagile\b", r"\bscrum\b", r"\bjira\b", r"\bkanban\b"],
    "System Design": [r"\bsystem\s+design\b", r"\bmicroservices\b", r"\bscalability\b", r"\bdistributed\s+systems\b"],
    "Testing / QA": [r"\bunit\s+testing\b", r"\bpytest\b", r"\bjest\b", r"\bcypress\b", r"\bselenium\b", r"\btdd\b"],
}

# 8 Comprehensive Tech Career Specialization Profiles with rich Curricula
CAREER_TRACKS = {
    "ai_ml": {
        "id": "ai_ml",
        "title": "AI & Machine Learning Engineer",
        "category": "Artificial Intelligence & Data Science",
        "tagline": "Build predictive models, neural networks, and cutting-edge GenAI/LLM solutions.",
        "icon": "brain",
        "badge_color": "purple",
        "difficulty": "Advanced",
        "demand": "Extremely High",
        "avg_salary_range": "$125,000 - $190,000 / yr",
        "target_roles": ["Machine Learning Engineer", "AI Solutions Architect", "NLP/LLM Engineer", "Applied Scientist"],
        "core_skills": [
            "Python", "Machine Learning", "Deep Learning", "PyTorch", "TensorFlow",
            "Generative AI & LLMs", "Natural Language Processing (NLP)", "Pandas", "NumPy", "SQL"
        ],
        "auxiliary_skills": [
            "Data Structures & Algorithms", "Computer Vision (CV)", "FastAPI", "Docker", "AWS", "Data Analysis / Visualization"
        ],
        "estimated_total_hours": 160,
        "phases": [
            {
                "phase_num": 1,
                "title": "Mathematical Foundations & Core Machine Learning",
                "duration_hours": 35,
                "focus": "Master Python data science stack, linear algebra, vector mathematics, and traditional ML algorithms.",
                "topics": [
                    "Advanced Python, NumPy vectorized computations & Pandas data manipulation",
                    "Exploratory Data Analysis (EDA), Feature Engineering & Imputation",
                    "Supervised Learning: Regression, Decision Trees, Random Forests, XGBoost",
                    "Unsupervised Learning: K-Means, PCA Dimensionality Reduction",
                    "Model Evaluation: Cross-validation, Precision/Recall, ROC-AUC, Bias-Variance tradeoff"
                ],
                "resources": [
                    {"name": "Andrew Ng - Machine Learning Specialization (Coursera)", "url": "https://www.coursera.org/specializations/machine-learning-introduction", "free": False},
                    {"name": "Fast.ai - Practical Deep Learning for Coders", "url": "https://course.fast.ai/", "free": True},
                    {"name": "Scikit-Learn Official User Guide & Tutorials", "url": "https://scikit-learn.org/stable/tutorial/index.html", "free": True}
                ],
                "milestone": "Build and evaluate an end-to-end customer churn/fraud detection pipeline with Scikit-Learn and XGBoost with >90% ROC-AUC."
            },
            {
                "phase_num": 2,
                "title": "Deep Learning & Modern Neural Architectures",
                "duration_hours": 45,
                "focus": "Deep dive into PyTorch, Backpropagation, CNNs, Transformers, and modern neural representations.",
                "topics": [
                    "PyTorch tensors, autograd engine, custom datasets and DataLoader pipelines",
                    "Multi-Layer Perceptrons, Regularization (Dropout, BatchNorm), Optimizers (AdamW)",
                    "Convolutional Neural Networks (CNNs) & Computer Vision fundamentals (ResNet, ViT)",
                    "Sequence Models & Self-Attention Mechanism: The Transformer Architecture in detail",
                    "Hugging Face Ecosystem: Tokenizers, Pre-trained Models, Model Hub"
                ],
                "resources": [
                    {"name": "DeepLearning.AI - Deep Learning Specialization", "url": "https://www.coursera.org/specializations/deep-learning", "free": False},
                    {"name": "Hugging Face NLP Course", "url": "https://huggingface.co/learn/nlp-course", "free": True},
                    {"name": "PyTorch Official 60-Minute Blitz & Tutorials", "url": "https://pytorch.org/tutorials/", "free": True}
                ],
                "milestone": "Implement a custom multi-head attention Transformer classifier in PyTorch from scratch on sentiment / classification datasets."
            },
            {
                "phase_num": 3,
                "title": "Generative AI, LLMs & Retrieval-Augmented Generation (RAG)",
                "duration_hours": 45,
                "focus": "Build production-grade GenAI applications, vector databases, LangChain/LlamaIndex, and LLM fine-tuning.",
                "topics": [
                    "Large Language Model Prompt Engineering, Structured Outputs & Tool Calling",
                    "Vector Embeddings, Similarity Search & Vector Databases (ChromaDB, Pinecone, FAISS)",
                    "Building Advanced RAG Pipelines: Chunking strategies, Re-ranking, HyDE, Multi-Query",
                    "Agentic AI: LangChain, LangGraph, autonomous agents with tool integration",
                    "Parameter Efficient Fine-Tuning (PEFT, LoRA, QLoRA) on open-weights models (Llama 3, Mistral)"
                ],
                "resources": [
                    {"name": "LangChain Official Conceptual & Production Guides", "url": "https://python.langchain.com/docs/get_started/introduction", "free": True},
                    {"name": "DeepLearning.AI - Building Systems with ChatGPT & LLMs", "url": "https://www.deeplearning.ai/short-courses/", "free": True},
                    {"name": "Pinecone Learning Center - Vector Search Mastery", "url": "https://www.pinecone.io/learn/", "free": True}
                ],
                "milestone": "Deploy a multi-document Enterprise Knowledge Assistant with RAG, source citation, hybrid search, and FastAPI endpoint."
            },
            {
                "phase_num": 4,
                "title": "MLOps, Model Deployment & Production Scale",
                "duration_hours": 35,
                "focus": "Deploy ML models as low-latency microservices with Docker, Kubernetes, monitoring, and CI/CD pipelines.",
                "topics": [
                    "Packaging models with FastAPI, ONNX Runtime, and TensorRT for low latency",
                    "Containerization with Docker & Multi-stage builds for GPU/CPU inference",
                    "Model tracking and registry using MLflow or Weights & Biases",
                    "CI/CD automated testing and drift monitoring in production (Evidently AI)",
                    "Cloud deployment on AWS SageMaker / GCP Vertex AI / Hugging Face Spaces"
                ],
                "resources": [
                    {"name": "Made With ML - MLOps Course by Goku Mohandas", "url": "https://madewithml.com/", "free": True},
                    {"name": "Full Stack Deep Learning Course", "url": "https://fullstackdeeplearning.com/", "free": True},
                    {"name": "AWS Certified Machine Learning - Specialty Prep", "url": "https://aws.amazon.com/certification/certified-machine-learning-specialty/", "free": False}
                ],
                "milestone": "Containerize your LLM/ML pipeline with Docker, establish automated MLflow logging, and deploy with streaming API responses."
            }
        ],
        "recommended_certifications": [
            "AWS Certified Machine Learning - Specialty (MLS-C01)",
            "Google Cloud Professional Machine Learning Engineer",
            "TensorFlow Developer Certificate (or DeepLearning.AI Specialization)"
        ]
    },

    "fullstack_dev": {
        "id": "fullstack_dev",
        "title": "Full-Stack Web & Cloud Application Engineer",
        "category": "Software Engineering",
        "tagline": "Architect modern web apps, high-throughput backend APIs, and slick reactive user interfaces.",
        "icon": "layers",
        "badge_color": "emerald",
        "difficulty": "Intermediate - Advanced",
        "demand": "Extremely High",
        "avg_salary_range": "$110,000 - $165,000 / yr",
        "target_roles": ["Full-Stack Software Engineer", "Backend Developer", "Frontend Architect", "Web Application Engineer"],
        "core_skills": [
            "Python", "Django", "JavaScript", "TypeScript", "React",
            "REST APIs", "PostgreSQL", "HTML/CSS", "Git & GitHub", "Docker"
        ],
        "auxiliary_skills": [
            "Next.js", "FastAPI", "Node.js", "Redis", "Tailwind CSS", "CI/CD", "AWS", "System Design"
        ],
        "estimated_total_hours": 140,
        "phases": [
            {
                "phase_num": 1,
                "title": "Modern Frontend Mastery (React, TypeScript & Responsive UI)",
                "duration_hours": 35,
                "focus": "Build responsive, accessible, component-driven client applications with state management.",
                "topics": [
                    "Modern JavaScript (ES6+), TypeScript interfaces, generics and strict typing",
                    "React 18+ architecture: Hooks (useState, useEffect, useMemo, useCallback, custom hooks)",
                    "State management with Context API, Zustand or Redux Toolkit",
                    "CSS Grid, Flexbox, Tailwind CSS and Glassmorphism design systems",
                    "Client-side routing, Form handling with React Hook Form & Zod schema validation"
                ],
                "resources": [
                    {"name": "React Official Interactive Documentation", "url": "https://react.dev/", "free": True},
                    {"name": "Full Stack Open (University of Helsinki)", "url": "https://fullstackopen.com/en/", "free": True},
                    {"name": "The Odin Project - Full Stack JavaScript", "url": "https://www.theodinproject.com/", "free": True}
                ],
                "milestone": "Create a fully responsive, type-safe interactive web dashboard with dynamic data tables, search, filtering, and theme switching."
            },
            {
                "phase_num": 2,
                "title": "Robust Backend Engineering (Django / FastAPI & PostgreSQL)",
                "duration_hours": 40,
                "focus": "Develop secure, scalable server architecture, ORM optimization, and RESTful/GraphQL APIs.",
                "topics": [
                    "Django ORM deep dive: query optimization, select_related/prefetch_related, indexing",
                    "Django REST Framework: Serializers, ViewSets, Token & JWT Authentication, Permissions",
                    "Relational Database modeling, ACID transactions, migrations, and PostgreSQL full-text search",
                    "Asynchronous task queues with Celery, Redis caching, and Background workers",
                    "API security: CSRF, CORS, Rate Limiting, SQL Injection & XSS sanitization"
                ],
                "resources": [
                    {"name": "Django Software Foundation Documentation", "url": "https://docs.djangoproject.com/", "free": True},
                    {"name": "TestDriven.io - Real World Django & React Microservices", "url": "https://testdriven.io/", "free": False},
                    {"name": "FastAPI Official Interactive Tutorial", "url": "https://fastapi.tiangolo.com/tutorial/", "free": True}
                ],
                "milestone": "Build a multi-tenant SaaS backend with JWT auth, role-based access control (RBAC), and automated PDF/CSV report generation via background jobs."
            },
            {
                "phase_num": 3,
                "title": "Full-Stack Integration, Real-Time WebSockets & Caching",
                "duration_hours": 35,
                "focus": "Connect decoupled frontend and backend systems with real-time sockets and distributed caching.",
                "topics": [
                    "Connecting React/Next.js frontend with Django/FastAPI backend via Axios / TanStack Query",
                    "Real-time bi-directional messaging with WebSockets (Django Channels / Socket.io)",
                    "In-memory caching patterns with Redis (Query caching, session store, rate limiters)",
                    "Server-Side Rendering (SSR) & Static Site Generation (SSG) with Next.js",
                    "File upload pipelines to AWS S3 / Cloud Storage with signed URLs"
                ],
                "resources": [
                    {"name": "TanStack Query (React Query) Documentation", "url": "https://tanstack.com/query/latest", "free": True},
                    {"name": "Next.js App Router Official Course", "url": "https://nextjs.org/learn", "free": True},
                    {"name": "Django Channels WebSockets Guide", "url": "https://channels.readthedocs.io/en/stable/", "free": True}
                ],
                "milestone": "Build an interactive real-time collaboration workspace (collaborative whiteboard or live chat with notifications and S3 media attachments)."
            },
            {
                "phase_num": 4,
                "title": "System Design, Containerization & Production CI/CD",
                "duration_hours": 30,
                "focus": "Containerize, test, and deploy full-stack applications with zero-downtime CI/CD workflows.",
                "topics": [
                    "Docker multi-stage builds for frontend and backend production containers",
                    "Docker Compose orchestration for multi-container local and staging environments",
                    "GitHub Actions CI/CD pipeline: Linting (Ruff/ESLint), automated PyTest/Jest, Docker image publishing",
                    "System Design fundamentals: Load balancing, CDN, Database read replicas, Horizontal scaling",
                    "Production monitoring: Sentry error tracking, Nginx reverse proxy, SSL/TLS configuration"
                ],
                "resources": [
                    {"name": "System Design Primer by Donne Martin (GitHub)", "url": "https://github.com/donnemartin/system-design-primer", "free": True},
                    {"name": "Docker & Kubernetes Documentation", "url": "https://docs.docker.com/get-started/", "free": True},
                    {"name": "GitHub Actions Documentation", "url": "https://docs.github.com/en/actions", "free": True}
                ],
                "milestone": "Orchestrate your full-stack application with Docker Compose, configure automated GitHub Actions testing, and deploy to a cloud instance."
            }
        ],
        "recommended_certifications": [
            "AWS Certified Developer - Associate (DVA-C02)",
            "Meta Certified Front-End / Back-End Developer",
            "PostgreSQL Certified Associate / Professional"
        ]
    },

    "devops_cloud": {
        "id": "devops_cloud",
        "title": "Cloud Solutions & DevOps Architect",
        "category": "Cloud & Infrastructure",
        "tagline": "Automate cloud infrastructure, manage Kubernetes clusters, and build bulletproof CI/CD pipelines.",
        "icon": "cloud",
        "badge_color": "sky",
        "difficulty": "Intermediate - Advanced",
        "demand": "Extremely High",
        "avg_salary_range": "$120,000 - $185,000 / yr",
        "target_roles": ["DevOps Engineer", "Cloud Solutions Architect", "Site Reliability Engineer (SRE)", "Platform Engineer"],
        "core_skills": [
            "Linux / Shell", "AWS", "Docker", "Kubernetes", "Terraform",
            "CI/CD", "Git & GitHub", "Python", "Prometheus / Grafana", "System Design"
        ],
        "auxiliary_skills": [
            "Google Cloud (GCP)", "Azure", "Go", "Network Security", "PostgreSQL", "Redis"
        ],
        "estimated_total_hours": 150,
        "phases": [
            {
                "phase_num": 1,
                "title": "Linux Systems Mastery, Networking & Containerization",
                "duration_hours": 35,
                "focus": "Master core Linux administration, network protocols, shell scripting, and container internals.",
                "topics": [
                    "Linux kernel basics, processes, memory management, systemd services, SSH hardening",
                    "Networking protocols: TCP/IP, DNS, HTTP/HTTPS, SSL/TLS, Subnetting, VPC concepts",
                    "Advanced Bash and Python scripting for infrastructure automation",
                    "Docker containerization: Layers, multi-stage builds, non-root users, security scanning (Trivy)",
                    "Docker Compose multi-service architecture & volume persistence"
                ],
                "resources": [
                    {"name": "Linux Foundation - Introduction to Linux (edX)", "url": "https://www.edx.org/course/introduction-to-linux", "free": True},
                    {"name": "Docker Official Deep Dive Guides", "url": "https://docs.docker.com/", "free": True},
                    {"name": "OverTheWire: Bandit (Linux Command-line Game)", "url": "https://overthewire.org/wargames/bandit/", "free": True}
                ],
                "milestone": "Write an automated Linux provisioning script that configures an Ubuntu server with secure firewall, Nginx reverse proxy, and Docker runtime."
            },
            {
                "phase_num": 2,
                "title": "Infrastructure as Code (Terraform) & Cloud Architecture (AWS)",
                "duration_hours": 40,
                "focus": "Design multi-region cloud infrastructure using declarative Terraform code.",
                "topics": [
                    "AWS core services: VPC, Subnets, Internet Gateways, NAT, EC2, ECS, S3, IAM policies & roles",
                    "Terraform HCL syntax: Providers, Resources, Variables, Outputs, State locking with S3 & DynamoDB",
                    "Terraform reusable modules, dynamic blocks, workspace environments (dev/stage/prod)",
                    "Cloud security best practices: Principle of least privilege, Secrets Manager, KMS encryption",
                    "Cost optimization, Auto-scaling groups, and Elastic Load Balancing (ALB/NLB)"
                ],
                "resources": [
                    {"name": "HashiCorp Learn - Terraform Associate Tutorials", "url": "https://developer.hashicorp.com/terraform/tutorials", "free": True},
                    {"name": "Stephane Maarek - AWS Certified Solutions Architect Associate (Udemy)", "url": "https://www.udemy.com/course/aws-certified-solutions-architect-associate-saa-c03/", "free": False},
                    {"name": "AWS Well-Architected Framework Whitepaper", "url": "https://aws.amazon.com/architecture/well-architected/", "free": True}
                ],
                "milestone": "Write a complete Terraform module that provisions a highly-available AWS VPC across 2 Availability Zones with ALB, ECS cluster, and RDS PostgreSQL."
            },
            {
                "phase_num": 3,
                "title": "Container Orchestration with Kubernetes (K8s)",
                "duration_hours": 45,
                "focus": "Deploy, scale, and manage resilient microservice applications on Kubernetes clusters.",
                "topics": [
                    "Kubernetes architecture: Control Plane (API server, etcd, scheduler), Worker Nodes (kubelet, kube-proxy)",
                    "K8s Core Objects: Pods, Deployments, ReplicaSets, StatefulSets, DaemonSets",
                    "Networking & Storage: Services (ClusterIP, NodePort, LoadBalancer), Ingress controllers, PVs, PVCs",
                    "ConfigMaps, Secrets, RBAC, Service Accounts, and NetworkPolicies",
                    "Helm package management: Creating custom charts, value templates, and dependency management",
                    "Managed K8s: AWS EKS / GCP GKE cluster provisioning and maintenance"
                ],
                "resources": [
                    {"name": "Kubernetes Official Interactive Documentation & Tutorials", "url": "https://kubernetes.io/docs/tutorials/", "free": True},
                    {"name": "Mumshad Mannambeth - CKA / CKAD Course (KodeKloud)", "url": "https://kodekloud.com/", "free": False},
                    {"name": "Kubernetes The Hard Way by Kelsey Hightower", "url": "https://github.com/kelseyhightower/kubernetes-the-hard-way", "free": True}
                ],
                "milestone": "Deploy a multi-service web application to a Kubernetes cluster using custom Helm charts, Ingress routing with TLS, and Horizontal Pod Autoscaling (HPA)."
            },
            {
                "phase_num": 4,
                "title": "Enterprise CI/CD, GitOps & Observability",
                "duration_hours": 30,
                "focus": "Automate zero-downtime release pipelines with GitOps (ArgoCD) and real-time observability.",
                "topics": [
                    "Advanced CI/CD with GitHub Actions: Matrix builds, container caching, automated canary releases",
                    "GitOps workflows with ArgoCD / Flux: Declarative application synchronization and rollbacks",
                    "Monitoring & Metrics: Prometheus metrics scraping, PromQL queries, Grafana dashboards",
                    "Logging & Distributed Tracing: OpenTelemetry, Grafana Loki / ELK Stack, Jaeger",
                    "Chaos engineering, SLA/SLO/SLI definition, and automated incident alerting (PagerDuty/Slack)"
                ],
                "resources": [
                    {"name": "ArgoCD Official GitOps Documentation", "url": "https://argo-cd.readthedocs.io/", "free": True},
                    {"name": "Prometheus & Grafana Official Guides", "url": "https://prometheus.io/docs/introduction/overview/", "free": True},
                    {"name": "Google Site Reliability Engineering (SRE) Books", "url": "https://sre.google/books/", "free": True}
                ],
                "milestone": "Implement a full GitOps pipeline where code commits trigger automated testing, container build, and continuous deployment to EKS via ArgoCD with Prometheus alerts."
            }
        ],
        "recommended_certifications": [
            "AWS Certified Solutions Architect - Associate (SAA-C03)",
            "Certified Kubernetes Administrator (CKA) / CKAD",
            "HashiCorp Certified: Terraform Associate (003)"
        ]
    },

    "data_engineering": {
        "id": "data_engineering",
        "title": "Data Engineering & Big Data Specialist",
        "category": "Data & Analytics",
        "tagline": "Build scalable data pipelines, distributed processing engines, and modern lakehouses.",
        "icon": "database",
        "badge_color": "amber",
        "difficulty": "Intermediate - Advanced",
        "demand": "Very High",
        "avg_salary_range": "$120,000 - $175,000 / yr",
        "target_roles": ["Data Engineer", "Big Data Engineer", "Analytics Engineer", "Data Platform Architect"],
        "core_skills": [
            "Python", "SQL", "Apache Spark", "Airflow", "PostgreSQL",
            "Data Warehousing", "Apache Kafka", "Docker", "Git & GitHub", "AWS"
        ],
        "auxiliary_skills": [
            "Pandas", "Linux / Shell", "Java", "Go", "Machine Learning", "Terraform"
        ],
        "estimated_total_hours": 150,
        "phases": [
            {
                "phase_num": 1,
                "title": "Advanced SQL, Data Modeling & Relational Architectures",
                "duration_hours": 35,
                "focus": "Master complex analytical SQL, dimensional modeling (Kimball), and relational database performance.",
                "topics": [
                    "Advanced SQL: Window functions, CTEs, Recursive queries, Partitioning, Execution plans (EXPLAIN)",
                    "Data Modeling: Star Schema, Snowflake Schema, Fact & Dimension tables, Slowly Changing Dimensions (SCD 1, 2, 3)",
                    "PostgreSQL internals: Indexing strategies (B-Tree, GIN, BRIN), Connection pooling, Vacuuming",
                    "Python for ETL: Optimized data extraction, chunking, and database ingestion with psycopg3 / SQLAlchemy"
                ],
                "resources": [
                    {"name": "The Data Warehouse Toolkit by Ralph Kimball", "url": "https://www.kimballgroup.com/data-warehouse-business-intelligence-resources/books/data-warehouse-dw-toolkit/", "free": False},
                    {"name": "Mode Analytics Advanced SQL Tutorial", "url": "https://mode.com/sql-tutorial/", "free": True},
                    {"name": "DataCamp - Data Engineering Career Track", "url": "https://www.datacamp.com/", "free": False}
                ],
                "milestone": "Design and build a normalized staging and star-schema analytical data warehouse in PostgreSQL with automated Python ingestion scripts."
            },
            {
                "phase_num": 2,
                "title": "Distributed Big Data Processing with Apache Spark & PySpark",
                "duration_hours": 45,
                "focus": "Process massive distributed datasets across clusters using Spark DataFrames and Structured Streaming.",
                "topics": [
                    "Apache Spark Architecture: Driver, Executors, Tasks, DAG Scheduler, Catalyst Optimizer",
                    "PySpark DataFrames & Spark SQL: Transformations vs Actions, Wide vs Narrow dependencies, Shuffling",
                    "Spark Performance Tuning: Partitioning, Broadcast joins, Caching strategies, Skew handling",
                    "Columnar file formats: Parquet, ORC, Snappy compression benefits",
                    "Delta Lake / Apache Iceberg: ACID transactions on Data Lakes, Time Travel, Schema enforcement"
                ],
                "resources": [
                    {"name": "Spark: The Definitive Guide by Bill Chambers & Matei Zaharia", "url": "https://www.oreilly.com/library/view/spark-the-definitive/9781491912201/", "free": False},
                    {"name": "PySpark Official Documentation & Guides", "url": "https://spark.apache.org/docs/latest/api/python/", "free": True},
                    {"name": "DataTalks.Club - Data Engineering Zoomcamp", "url": "https://github.com/DataTalksClub/data-engineering-zoomcamp", "free": True}
                ],
                "milestone": "Process 10+ GB of raw log data using PySpark on a distributed cluster, transforming it into optimized Parquet tables on Delta Lake."
            },
            {
                "phase_num": 3,
                "title": "Workflow Orchestration (Airflow) & Cloud Data Warehousing (Snowflake/BigQuery)",
                "duration_hours": 40,
                "focus": "Orchestrate complex DAG data pipelines with Airflow and transform data with modern DBT.",
                "topics": [
                    "Apache Airflow fundamentals: DAGs, Operators, Sensors, Taskflow API, Hooks & Connections",
                    "Airflow scheduling, backfilling, retries, SLAs, dynamic task mapping",
                    "Modern Cloud Data Warehouses: Snowflake architecture, BigQuery slot management",
                    "dbt (data build tool): Data modeling, testing, documentation, Jinja macros, CI/CD for analytics"
                ],
                "resources": [
                    {"name": "Astronomer - Apache Airflow Certification Guides", "url": "https://www.astronomer.io/academy/", "free": True},
                    {"name": "dbt Fundamentals Official Course", "url": "https://courses.getdbt.com/courses/fundamentals", "free": True},
                    {"name": "Snowflake Official Hands-On Essentials", "url": "https://learn.snowflake.com/", "free": True}
                ],
                "milestone": "Build an end-to-end automated ELT pipeline in Airflow that pulls external API data, loads into Snowflake/BigQuery, and transforms models using dbt with automated tests."
            },
            {
                "phase_num": 4,
                "title": "Real-Time Streaming with Kafka & Production Infrastructure",
                "duration_hours": 30,
                "focus": "Build event-driven real-time stream processing pipelines with Apache Kafka.",
                "topics": [
                    "Apache Kafka architecture: Brokers, Topics, Partitions, Consumer Groups, Offsets, Replication",
                    "Kafka Producers & Consumers with Python/Java: Idempotence, schema registry (Avro/Protobuf)",
                    "Real-time stream processing with Spark Structured Streaming / Flink",
                    "Data quality & monitoring: Great Expectations, Monte Carlo data observability concepts"
                ],
                "resources": [
                    {"name": "Confluent Developer - Apache Kafka 101 Courses", "url": "https://developer.confluent.io/courses/", "free": True},
                    {"name": "Great Expectations Official Quickstart", "url": "https://docs.greatexpectations.io/", "free": True}
                ],
                "milestone": "Implement a real-time event pipeline where simulated user clickstreams are published to Kafka, processed with Spark Streaming, and stored in an analytical sink."
            }
        ],
        "recommended_certifications": [
            "AWS Certified Data Engineer - Associate (DEA-C01)",
            "Databricks Certified Data Engineer Associate / Professional",
            "dbt Certified Developer"
        ]
    },

    "cybersecurity": {
        "id": "cybersecurity",
        "title": "Cybersecurity & Ethical Hacking Specialist",
        "category": "Security & Defense",
        "tagline": "Protect infrastructure, conduct penetration testing, and analyze vulnerabilities.",
        "icon": "shield",
        "badge_color": "rose",
        "difficulty": "Intermediate - Advanced",
        "demand": "Extremely High",
        "avg_salary_range": "$115,000 - $170,000 / yr",
        "target_roles": ["Penetration Tester", "SOC Analyst", "Security Engineer", "Information Security Specialist"],
        "core_skills": [
            "Linux / Shell", "Network Security", "Penetration Testing", "Cryptography & Identity",
            "SIEM & SOC Analysis", "Python", "Git & GitHub", "AWS"
        ],
        "auxiliary_skills": [
            "Docker", "REST APIs", "C++", "Go", "SQL", "System Design"
        ],
        "estimated_total_hours": 140,
        "phases": [
            {
                "phase_num": 1,
                "title": "Networking Protocols, System Architecture & Linux Security",
                "duration_hours": 35,
                "focus": "Master OSI model, network traffic analysis with Wireshark, and Linux system hardening.",
                "topics": [
                    "TCP/IP handshakes, ARP, DNS spoofing, ICMP, routing protocols, subnets & VLANs",
                    "Packet inspection with Wireshark and tcpdump: Analyzing malformed traffic and attack signatures",
                    "Linux security: Permissions, SUID binaries, iptables/UFW, AppArmor, SELinux, auditd",
                    "Python for security: Socket programming, port scanners, banner grabbing, and packet crafting (Scapy)"
                ],
                "resources": [
                    {"name": "TryHackMe - Pre-Security & Complete Beginner Paths", "url": "https://tryhackme.com/", "free": True},
                    {"name": "Professor Messer's CompTIA Security+ (SY0-701) Training", "url": "https://www.professormesser.com/", "free": True},
                    {"name": "Wireshark Official User Guide", "url": "https://www.wireshark.org/docs/", "free": True}
                ],
                "milestone": "Capture and analyze network captures of various attacks (SYN flood, ARP poison, brute-force) in Wireshark and write a Python packet analyzer."
            },
            {
                "phase_num": 2,
                "title": "Web Application Security & OWASP Top 10",
                "duration_hours": 40,
                "focus": "Discover and exploit web application vulnerabilities ethically using industry tools.",
                "topics": [
                    "OWASP Top 10 vulnerabilities: SQL Injection, XSS, CSRF, SSRF, IDOR, Broken Authentication",
                    "Using Burp Suite Professional / Community for request interception, repeater, and intruder attacks",
                    "Authentication security: JWT vulnerabilities, OAuth 2.0 misconfigurations, Session fixation",
                    "API penetration testing: Insecure endpoints, mass assignment, rate-limit bypassing"
                ],
                "resources": [
                    {"name": "PortSwigger Web Security Academy (Burp Suite Labs)", "url": "https://portswigger.net/web-security", "free": True},
                    {"name": "OWASP Official Documentation & Cheatsheets", "url": "https://cheatsheetseries.owasp.org/", "free": True},
                    {"name": "Hack The Box Academy - Bug Bounty Hunter Path", "url": "https://academy.hackthebox.com/", "free": False}
                ],
                "milestone": "Complete 25+ PortSwigger Web Security Academy labs covering SQLi, XSS, SSRF, and authentication bypasses."
            },
            {
                "phase_num": 3,
                "title": "Network Penetration Testing & Privilege Escalation",
                "duration_hours": 35,
                "focus": "Perform full offensive assessments, active directory enumeration, and privilege escalation.",
                "topics": [
                    "Reconnaissance and enumeration with Nmap, Gobuster, Nikto, and OSINT techniques",
                    "Exploitation frameworks: Metasploit, searchsploit, custom exploit modification",
                    "Linux privilege escalation: Vulnerable cronjobs, SUID abuse, kernel exploits, sudo misconfigurations",
                    "Windows privilege escalation & Active Directory basics (Kerberoasting, Pass-the-Hash)"
                ],
                "resources": [
                    {"name": "Hack The Box - Active Labs & Pro Labs", "url": "https://www.hackthebox.com/", "free": False},
                    {"name": "TCM Security - Practical Ethical Hacking (PEH)", "url": "https://academy.tcm-sec.com/", "free": False},
                    {"name": "PayloadsAllTheThings (GitHub Offensive Cheatsheet)", "url": "https://github.com/swisskyrepo/PayloadsAllTheThings", "free": True}
                ],
                "milestone": "Successfully compromise and obtain root/system on 10 intermediate target machines on HackTheBox or VulnHub with detailed writeups."
            },
            {
                "phase_num": 4,
                "title": "Defensive Security (Blue Team), SIEM & Incident Response",
                "duration_hours": 30,
                "focus": "Investigate security incidents, build detection rules in SIEMs, and harden cloud environments.",
                "topics": [
                    "SIEM platforms (Splunk, Elastic SIEM, Microsoft Sentinel): Ingesting logs and creating alert queries",
                    "MITRE ATT&CK Framework: Mapping adversary tactics, techniques, and procedures (TTPs)",
                    "Threat hunting, memory analysis (Volatility), and endpoint detection & response (EDR)",
                    "Cloud Security: AWS IAM hardening, GuardDuty, CloudTrail analysis, Security Hub"
                ],
                "resources": [
                    {"name": "Splunk Free Training & Search Tutorial", "url": "https://www.splunk.com/en_us/training/free-courses/splunk-fundamentals-1.html", "free": True},
                    {"name": "MITRE ATT&CK Matrix Navigator", "url": "https://attack.mitre.org/", "free": True},
                    {"name": "CyberDefenders - Blue Team Training Labs", "url": "https://cyberdefenders.org/", "free": True}
                ],
                "milestone": "Configure a Splunk SIEM instance, simulate a web attack, create custom detection alerts, and generate a formal Incident Response Report."
            }
        ],
        "recommended_certifications": [
            "CompTIA Security+ (SY0-701)",
            "Offensive Security Certified Professional (OSCP) or eJPT",
            "Certified Information Systems Security Professional (CISSP - for senior)"
        ]
    },

    "mobile_dev": {
        "id": "mobile_dev",
        "title": "Mobile App Developer (Cross-Platform / Native)",
        "category": "Mobile Software Engineering",
        "tagline": "Craft sleek, high-performance iOS and Android mobile apps with Flutter or React Native.",
        "icon": "smartphone",
        "badge_color": "cyan",
        "difficulty": "Intermediate",
        "demand": "High",
        "avg_salary_range": "$105,000 - $155,000 / yr",
        "target_roles": ["Mobile Application Engineer", "Flutter Developer", "React Native Engineer", "iOS/Android Developer"],
        "core_skills": [
            "Flutter", "React Native", "JavaScript", "TypeScript", "Kotlin",
            "Swift", "REST APIs", "Git & GitHub", "Firebase"
        ],
        "auxiliary_skills": [
            "Python", "UI/UX Design", "Figma", "Docker", "Testing / QA"
        ],
        "estimated_total_hours": 130,
        "phases": [
            {
                "phase_num": 1,
                "title": "Core Mobile Frameworks (Flutter / React Native) & State",
                "duration_hours": 35,
                "focus": "Build responsive mobile user interfaces and state management architectures.",
                "topics": [
                    "Dart / modern JavaScript fundamentals, async programming, streams & futures",
                    "Widget tree / Component hierarchy, responsive layouts, theme switching",
                    "State management: Riverpod / Bloc for Flutter, Zustand / Redux for React Native",
                    "Mobile navigation patterns, deep linking, tab bars, drawer menus"
                ],
                "resources": [
                    {"name": "Flutter Official Documentation & Codelabs", "url": "https://flutter.dev/docs/codelabs", "free": True},
                    {"name": "React Native Official Tutorial", "url": "https://reactnative.dev/docs/getting-started", "free": True}
                ],
                "milestone": "Build a multi-screen mobile ecommerce catalog app with smooth transition animations and state persistence."
            },
            {
                "phase_num": 2,
                "title": "Device APIs, Local Storage & Backend Integration",
                "duration_hours": 35,
                "focus": "Integrate camera, GPS geolocation, push notifications, and offline caching.",
                "topics": [
                    "Accessing native device features: Camera, Photo library, Geolocation, Biometrics",
                    "Offline-first architecture: SQLite, Hive, WatermelonDB, or Realm database synchronization",
                    "Networking with Dio / Axios: Interceptors, token refreshing, retry mechanisms",
                    "Push notifications setup with Firebase Cloud Messaging (FCM) / OneSignal"
                ],
                "resources": [
                    {"name": "Firebase for Mobile Documentation", "url": "https://firebase.google.com/docs", "free": True},
                    {"name": "Reso Coder Flutter Clean Architecture Tutorial", "url": "https://resocoder.com/flutter-clean-architecture-tdd/", "free": True}
                ],
                "milestone": "Create an offline-first fitness or travel tracker app with GPS mapping, image uploads, and background sync."
            },
            {
                "phase_num": 3,
                "title": "Performance Optimization, Animations & Security",
                "duration_hours": 30,
                "focus": "Optimize render frames to 60/120 FPS, handle memory management, and secure credentials.",
                "topics": [
                    "Frame budget, avoiding UI thread jank, optimizing heavy list rendering (ListView.builder / FlashList)",
                    "Custom animations with Flutter Rive / React Native Reanimated",
                    "Mobile security: Secure storage (KeyStore / Keychain), certificate pinning, obfuscation with ProGuard"
                ],
                "resources": [
                    {"name": "React Native Performance Optimization Guide", "url": "https://reactnative.dev/docs/profile-hermes", "free": True},
                    {"name": "Flutter Performance Profiling & DevTools", "url": "https://docs.flutter.dev/perf", "free": True}
                ],
                "milestone": "Build an ultra-smooth interactive social feed with micro-interactions, gesture-driven card swipes, and 60fps animations."
            },
            {
                "phase_num": 4,
                "title": "App Store Deployment & CI/CD Pipelines (Fastlane)",
                "duration_hours": 30,
                "focus": "Automate beta testing distributions (TestFlight / Google Play Beta) and store publishing.",
                "topics": [
                    "App signing, provisioning profiles, keystore management for iOS & Android",
                    "Automating builds and store submissions with Fastlane and GitHub Actions",
                    "Crashlytics crash reporting and in-app analytics integration",
                    "In-App Purchases (RevenueCat / StoreKit) and App Store Review guidelines compliance"
                ],
                "resources": [
                    {"name": "Fastlane Official Mobile CI/CD Documentation", "url": "https://docs.fastlane.tools/", "free": True},
                    {"name": "Google Play & Apple Developer Publishing Guides", "url": "https://developer.apple.com/distribute/", "free": True}
                ],
                "milestone": "Configure automated Fastlane scripts to build, sign, and push internal test builds to Google Play Internal Track and TestFlight."
            }
        ],
        "recommended_certifications": [
            "Google Associate Android Developer Certification",
            "Meta Certified iOS / Android Developer",
            "Flutter Certified Application Developer"
        ]
    },

    "frontend_uiux": {
        "id": "frontend_uiux",
        "title": "Frontend Engineer & UI/UX Specialist",
        "category": "Frontend & Design Engineering",
        "tagline": "Design high-fidelity user experiences and code pixel-perfect modern web interfaces.",
        "icon": "layout",
        "badge_color": "teal",
        "difficulty": "Intermediate",
        "demand": "High",
        "avg_salary_range": "$105,000 - $155,000 / yr",
        "target_roles": ["Frontend Engineer", "UI/UX Developer", "Design Technologist", "Design System Engineer"],
        "core_skills": [
            "HTML/CSS", "JavaScript", "TypeScript", "React", "Next.js",
            "Tailwind CSS", "Figma", "UI/UX Design", "Git & GitHub"
        ],
        "auxiliary_skills": [
            "Vue.js", "REST APIs", "Testing / QA", "Node.js", "Python"
        ],
        "estimated_total_hours": 130,
        "phases": [
            {
                "phase_num": 1,
                "title": "UI/UX Foundations, Design Systems & Figma Prototyping",
                "duration_hours": 30,
                "focus": "Master typography, color theory, spacing tokens, accessibility, and interactive Figma prototyping.",
                "topics": [
                    "Design principles: Visual hierarchy, contrast, white space, micro-typography",
                    "Figma mastery: Auto-layout, components, variants, variables, design tokens",
                    "Web Accessibility (a11y): WCAG 2.2 AA standards, screen readers, ARIA roles, keyboard navigation",
                    "Building a consistent Design System (tokens for spacing, color palettes, typography scales)"
                ],
                "resources": [
                    {"name": "Refactoring UI by Adam Wathan & Steve Schoger", "url": "https://www.refactoringui.com/", "free": False},
                    {"name": "Figma Official Design Systems Interactive Course", "url": "https://www.figma.com/resources/learn-design/", "free": True},
                    {"name": "Web.dev Accessibility Guide by Google", "url": "https://web.dev/learn/accessibility/", "free": True}
                ],
                "milestone": "Design a complete SaaS web application in Figma with components, tokens, light/dark modes, and interactive click-through prototype."
            },
            {
                "phase_num": 2,
                "title": "Modern Component Engineering (React, Next.js & Tailwind)",
                "duration_hours": 40,
                "focus": "Translate Figma designs into robust, reusable, accessible React components with Tailwind CSS.",
                "topics": [
                    "Modern component architectures (Headless UI, Radix UI, Shadcn/ui patterns)",
                    "Tailwind CSS advanced utilities, custom plugins, container queries, and animations",
                    "State machines and complex form handling (React Hook Form, Zod, TanStack Table)",
                    "Next.js App Router: Server vs Client Components, Streaming UI with Suspense"
                ],
                "resources": [
                    {"name": "Shadcn UI Architecture & Source Code Study", "url": "https://ui.shadcn.com/", "free": True},
                    {"name": "Tailwind CSS Official Documentation", "url": "https://tailwindcss.com/docs", "free": True},
                    {"name": "Frontend Masters - Design for Developers", "url": "https://frontendmasters.com/", "free": False}
                ],
                "milestone": "Build an open-source React component library with Storybook documentation and automated visual regression testing."
            },
            {
                "phase_num": 3,
                "title": "Advanced Web Animation, Micro-Interactions & 3D Web",
                "duration_hours": 35,
                "focus": "Create awe-inspiring micro-interactions, scroll-driven animations, and interactive 3D elements.",
                "topics": [
                    "Framer Motion: Layout animations, spring physics, drag gestures, shared element transitions",
                    "CSS Scroll-driven animations, View Transitions API, and CSS GPU hardware acceleration",
                    "Introduction to 3D on the Web: Three.js, React Three Fiber (R3F), and Spline 3D embeds",
                    "Audio/haptic feedback on the web and progressive enhancement"
                ],
                "resources": [
                    {"name": "Framer Motion Official Interactive Guide", "url": "https://www.framer.com/motion/", "free": True},
                    {"name": "Three.js Journey by Bruno Simon", "url": "https://threejs-journey.com/", "free": False},
                    {"name": "Codrops Web Design & Experimental Interaction Tutorials", "url": "https://tympanus.net/codrops/", "free": True}
                ],
                "milestone": "Build a stunning award-winning landing page with 3D interactive hero canvas, smooth scroll reveal animations, and glassmorphism styling."
            },
            {
                "phase_num": 4,
                "title": "Web Performance (Core Web Vitals) & Production Testing",
                "duration_hours": 25,
                "focus": "Optimize Largest Contentful Paint (LCP), Interaction to Next Paint (INP), and test with Playwright.",
                "topics": [
                    "Core Web Vitals optimization: LCP, INP, CLS, font subsetting, image optimization, dynamic imports",
                    "Bundle analysis, code splitting, tree shaking, and modern bundlers (Vite / Turbopack)",
                    "Testing with Jest, React Testing Library, and End-to-End browser tests with Playwright / Cypress"
                ],
                "resources": [
                    {"name": "Google Web.dev Core Web Vitals Guides", "url": "https://web.dev/explore/fast", "free": True},
                    {"name": "Playwright Official Testing Documentation", "url": "https://playwright.dev/", "free": True}
                ],
                "milestone": "Achieve 100/100 Lighthouse Performance and Accessibility scores on a dynamic Next.js application with full E2E Playwright coverage."
            }
        ],
        "recommended_certifications": [
            "Meta Front-End Developer Professional Certificate",
            "W3Cx Front-End Web Developer Professional Certificate",
            "Nielsen Norman Group UX Master Certified"
        ]
    },

    "data_analytics": {
        "id": "data_analytics",
        "title": "Data Analyst & Business Intelligence Specialist",
        "category": "Data & Analytics",
        "tagline": "Turn raw business data into actionable strategic insights, executive dashboards, and statistical models.",
        "icon": "bar-chart",
        "badge_color": "indigo",
        "difficulty": "Beginner - Intermediate",
        "demand": "High",
        "avg_salary_range": "$85,000 - $130,000 / yr",
        "target_roles": ["Data Analyst", "Business Intelligence Analyst", "Product Analyst", "Operations Data Specialist"],
        "core_skills": [
            "SQL", "Python", "Data Analysis / Visualization", "Pandas", "NumPy",
            "PostgreSQL", "Git & GitHub"
        ],
        "auxiliary_skills": [
            "Machine Learning", "Data Warehousing", "HTML/CSS", "Django", "FastAPI"
        ],
        "estimated_total_hours": 120,
        "phases": [
            {
                "phase_num": 1,
                "title": "SQL for Analytics & Relational Databases",
                "duration_hours": 30,
                "focus": "Master SQL querying, aggregation, multi-table joins, and cohort analysis.",
                "topics": [
                    "SQL fundamentals: SELECT, WHERE, GROUP BY, HAVING, complex INNER/LEFT/FULL joins",
                    "Analytical SQL: Window functions (ROW_NUMBER, RANK, DENSE_RANK, LEAD, LAG), CTEs",
                    "Cohort retention analysis, user churn rate calculations, month-over-month growth metrics",
                    "Database optimization: Query plans, subquery vs join performance"
                ],
                "resources": [
                    {"name": "Mode Analytics SQL Tutorial", "url": "https://mode.com/sql-tutorial/", "free": True},
                    {"name": "Stratascratch Real Interview SQL Questions", "url": "https://www.stratascratch.com/", "free": True}
                ],
                "milestone": "Analyze an ecommerce dataset with 500k+ rows in PostgreSQL to identify customer segments, purchase churn patterns, and CLV (Customer Lifetime Value)."
            },
            {
                "phase_num": 2,
                "title": "Python for Data Analysis (Pandas, NumPy & Seaborn)",
                "duration_hours": 35,
                "focus": "Clean, manipulate, and visualize complex datasets programmatically.",
                "topics": [
                    "Data wrangling with Pandas: Handling missing data, pivoting, melting, group operations, datetime parsing",
                    "NumPy vectorized operations and statistical calculations",
                    "Exploratory Data Analysis (EDA) methodologies and hypothesis generation",
                    "Data visualization with Matplotlib, Seaborn, and interactive Plotly charts"
                ],
                "resources": [
                    {"name": "Python for Data Analysis by Wes McKinney (O'Reilly)", "url": "https://wesmckinney.com/book/", "free": True},
                    {"name": "Kaggle Micro-Courses for Python & Pandas", "url": "https://www.kaggle.com/learn", "free": True}
                ],
                "milestone": "Produce an in-depth exploratory data analysis report in Jupyter Notebook with interactive Plotly visualizations diagnosing business bottlenecks."
            },
            {
                "phase_num": 3,
                "title": "Business Intelligence & Executive Dashboards (Tableau / Power BI)",
                "duration_hours": 30,
                "focus": "Build automated executive BI dashboards, KPI tracking metrics, and data storytelling.",
                "topics": [
                    "Data modeling in Power BI / Tableau: Star schema, relationships, calculated fields / DAX",
                    "Dashboard visual layout: Designing for executive scanning, KPI summary cards, drill-through filters",
                    "Data storytelling: Translating analytical findings into actionable business recommendations",
                    "Scheduled automated data refreshes and alert triggers"
                ],
                "resources": [
                    {"name": "Microsoft Power BI Official Guided Learning", "url": "https://learn.microsoft.com/en-us/power-bi/guided-learning/", "free": True},
                    {"name": "Tableau Public Community & Tutorials", "url": "https://www.tableau.com/learn/training", "free": True}
                ],
                "milestone": "Build an executive-ready multi-tab interactive BI dashboard in Power BI/Tableau tracking sales velocity, regional performance, and forecasted targets."
            },
            {
                "phase_num": 4,
                "title": "Applied Statistics & A/B Testing for Product Decisions",
                "duration_hours": 25,
                "focus": "Design and analyze randomized controlled experiments (A/B tests) and basic predictive modeling.",
                "topics": [
                    "Probability distributions, Central Limit Theorem, confidence intervals, p-values",
                    "A/B test design: Sample size determination, statistical power, Type I & II errors, Bonferroni correction",
                    "Linear and Logistic Regression for business trend forecasting with statsmodels / scikit-learn",
                    "Presenting findings to stakeholders with structured executive summaries"
                ],
                "resources": [
                    {"name": "Khan Academy - Statistics and Probability", "url": "https://www.khanacademy.org/math/statistics-probability", "free": True},
                    {"name": "Udacity - A/B Testing by Google", "url": "https://www.udacity.com/course/ab-testing--ud257", "free": True}
                ],
                "milestone": "Design an A/B testing experiment for a landing page conversion redesign, analyze results in Python with statistical significance tests, and write a decision memo."
            }
        ],
        "recommended_certifications": [
            "Google Data Analytics Professional Certificate",
            "Microsoft Certified: Power BI Data Analyst Associate (PL-300)",
            "Tableau Certified Data Analyst"
        ]
    }
}


def extract_text_from_file(file_obj, filename: str) -> str:
    """
    Extracts text content from uploaded file (PDF, DOCX, TXT).
    Handles encoding nuances and error fallbacks.
    """
    ext = filename.lower().split('.')[-1]
    
    # Read binary bytes
    content_bytes = file_obj.read()
    if hasattr(file_obj, 'seek'):
        file_obj.seek(0)
        
    text = ""
    
    if ext == "pdf":
        if pypdf:
            try:
                reader = pypdf.PdfReader(io.BytesIO(content_bytes))
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
            except Exception as e:
                text = f"[PDF extraction note: Error reading binary content: {str(e)}]"
        else:
            text = "[PDF parser pypdf not available]"
            
    elif ext in ["docx", "doc"]:
        if docx:
            try:
                doc = docx.Document(io.BytesIO(content_bytes))
                for para in doc.paragraphs:
                    text += para.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                        text += "\n"
            except Exception as e:
                # Fallback to UTF-8 / latin-1 decoding
                text = content_bytes.decode("utf-8", errors="ignore")
        else:
            text = content_bytes.decode("utf-8", errors="ignore")
            
    else:
        # Default text file decoding
        try:
            text = content_bytes.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text = content_bytes.decode("latin-1")
            except Exception:
                text = content_bytes.decode("ascii", errors="ignore")
                
    return text.strip()


def extract_skills_from_text(raw_text: str) -> List[str]:
    """
    Scans raw resume text against the skill taxonomy and returns a list of detected skills.
    """
    detected_skills = []
    text_lower = raw_text.lower()
    
    for skill_name, patterns in SKILL_PATTERNS.items():
        for pattern in patterns:
            if re.search(pattern, text_lower, re.IGNORECASE):
                detected_skills.append(skill_name)
                break
                
    return detected_skills


def estimate_experience_level(raw_text: str) -> Dict[str, Any]:
    """
    Estimates candidate's experience level (Junior, Mid, Senior, Lead) from years of experience and keywords.
    """
    text_lower = raw_text.lower()
    
    # Check for senior / lead titles
    is_lead = bool(re.search(r"\b(lead|principal|staff|architect|director|head of)\b", text_lower))
    is_senior = bool(re.search(r"\b(senior|sr\.?|5\+\s*years|6\+\s*years|7\+\s*years|8\+\s*years)\b", text_lower))
    is_junior = bool(re.search(r"\b(intern|internship|junior|jr\.?|student|graduate|entry[- ]level|fresher)\b", text_lower))
    
    # Try finding years of experience
    years_matches = re.findall(r"(\d{1,2})\+?\s*(?:years?|yrs?)(?:\s+of)?\s+(?:experience|exp)", text_lower)
    max_years = 0
    if years_matches:
        try:
            max_years = max([int(y) for y in years_matches if int(y) < 40])
        except Exception:
            max_years = 0
            
    if is_lead or max_years >= 8:
        level = "Senior / Tech Lead"
        badge = "Senior"
    elif is_senior or max_years >= 4:
        level = "Mid-to-Senior Professional"
        badge = "Mid-Senior"
    elif is_junior or max_years <= 1:
        level = "Early Career / Aspiring Specialist"
        badge = "Entry Level"
    else:
        level = "Intermediate Practitioner"
        badge = "Intermediate"
        
    return {
        "level": level,
        "badge": badge,
        "detected_years": max_years if max_years > 0 else "1-3 (Estimated)"
    }


def analyze_resume_and_recommend_pathway(
    raw_text: str,
    weekly_hours: int = 10,
    preferred_track: Optional[str] = None
) -> Dict[str, Any]:
    """
    Core Intelligence Engine:
    1. Extracts acquired skills.
    2. Computes compatibility and match score with each career track.
    3. Identifies skill gaps and strengths.
    4. Selects the top recommended pathway (or user-chosen preference) and generates dynamic roadmap.
    """
    detected_skills = extract_skills_from_text(raw_text)
    exp_info = estimate_experience_level(raw_text)
    
    scored_tracks = []
    
    for track_key, track_data in CAREER_TRACKS.items():
        core_skills = track_data["core_skills"]
        aux_skills = track_data["auxiliary_skills"]
        
        # Calculate skills acquired vs missing
        acquired_core = [s for s in core_skills if s in detected_skills]
        missing_core = [s for s in core_skills if s not in detected_skills]
        
        acquired_aux = [s for s in aux_skills if s in detected_skills]
        missing_aux = [s for s in aux_skills if s not in detected_skills]
        
        # Weighted score: Core skills 75%, Auxiliary skills 25%
        core_pct = len(acquired_core) / max(len(core_skills), 1)
        aux_pct = len(acquired_aux) / max(len(aux_skills), 1)
        
        raw_match_score = int(round((core_pct * 75) + (aux_pct * 25)))
        # Normalize between 18% and 98% for realistic feel
        display_score = max(min(raw_match_score, 98), 18) if len(detected_skills) > 0 else 20
        
        scored_tracks.append({
            "id": track_key,
            "title": track_data["title"],
            "category": track_data["category"],
            "tagline": track_data["tagline"],
            "icon": track_data["icon"],
            "badge_color": track_data["badge_color"],
            "match_score": display_score,
            "core_match_count": len(acquired_core),
            "core_total_count": len(core_skills),
            "acquired_core": acquired_core,
            "missing_core": missing_core,
            "acquired_aux": acquired_aux,
            "missing_aux": missing_aux,
            "all_acquired": acquired_core + acquired_aux,
            "all_missing": missing_core + missing_aux,
            "difficulty": track_data["difficulty"],
            "demand": track_data["demand"],
            "avg_salary_range": track_data["avg_salary_range"],
            "target_roles": track_data["target_roles"],
            "estimated_total_hours": track_data["estimated_total_hours"],
            "phases": track_data["phases"],
            "recommended_certifications": track_data["recommended_certifications"]
        })
        
    # Sort tracks by match score descending
    scored_tracks.sort(key=lambda x: x["match_score"], reverse=True)
    
    # Pick primary track (either preferred by user or top scoring)
    chosen_track = None
    if preferred_track and preferred_track in CAREER_TRACKS:
        for t in scored_tracks:
            if t["id"] == preferred_track:
                chosen_track = t
                break
                
    if not chosen_track:
        chosen_track = scored_tracks[0]
        
    # Compute weekly pace and timeline dates
    weekly_hours = max(int(weekly_hours), 5)
    total_hours = chosen_track["estimated_total_hours"]
    total_weeks = max(int(round(total_hours / weekly_hours)), 4)
    
    # Build personalized phased schedule
    start_date = datetime.now()
    phased_roadmap = []
    accumulated_hours = 0
    
    for phase in chosen_track["phases"]:
        phase_hours = phase["duration_hours"]
        phase_weeks = max(int(round(phase_hours / weekly_hours)), 1)
        phase_start = start_date + timedelta(weeks=int(round(accumulated_hours / weekly_hours)))
        phase_end = phase_start + timedelta(weeks=phase_weeks)
        accumulated_hours += phase_hours
        
        phased_roadmap.append({
            "phase_num": phase["phase_num"],
            "title": phase["title"],
            "duration_hours": phase_hours,
            "estimated_weeks": phase_weeks,
            "start_date": phase_start.strftime("%b %d, %Y"),
            "end_date": phase_end.strftime("%b %d, %Y"),
            "focus": phase["focus"],
            "topics": phase["topics"],
            "resources": phase["resources"],
            "milestone": phase["milestone"]
        })
        
    completion_date = start_date + timedelta(weeks=total_weeks)
    
    return {
        "candidate_experience": exp_info,
        "detected_skills_count": len(detected_skills),
        "detected_skills": detected_skills,
        "primary_track": chosen_track,
        "alternative_tracks": [t for t in scored_tracks if t["id"] != chosen_track["id"]],
        "weekly_hours": weekly_hours,
        "total_estimated_hours": total_hours,
        "total_estimated_weeks": total_weeks,
        "completion_date": completion_date.strftime("%B %d, %Y"),
        "roadmap": phased_roadmap,
        "raw_text_length": len(raw_text),
        "raw_text_snippet": raw_text[:300] + ("..." if len(raw_text) > 300 else "")
    }


def get_sample_resumes() -> Dict[str, Dict[str, str]]:
    """
    Returns pre-packaged realistic sample resumes for instant testing.
    """
    return {
        "junior_python_web": {
            "title": "Junior Python / Django Developer",
            "desc": "Proficient in Python, Django, basic HTML/CSS, SQL, Git with interest in Full-Stack engineering.",
            "text": """Alex Morgan
Email: alex.morgan.dev@example.com | GitHub: github.com/alexmorgan | LinkedIn: linkedin.com/in/alexmorgan
Location: San Francisco, CA

SUMMARY
Junior Software Engineer with 2 years of academic and freelance experience building web applications using Python and Django. Passionate about clean code, RESTful API design, and modern database management with PostgreSQL. Seeking a Full-Stack or Backend Engineering trajectory.

TECHNICAL SKILLS
- Languages: Python, JavaScript, HTML5, CSS3, SQL
- Frameworks & Tools: Django, Django REST Framework, Bootstrap, Git, GitHub, SQLite, PostgreSQL
- Concepts: Object-Oriented Programming (OOP), Data Structures & Algorithms, REST APIs, Agile / Scrum

PROJECTS
1. E-Commerce Bookstore API (Django, PostgreSQL, Stripe)
- Designed and built a modular REST API using Django and Django REST Framework.
- Implemented user authentication with JWT tokens and payment processing integration.
- Managed database schemas and migrations with PostgreSQL.

2. Task Management Web App (Django, JavaScript, HTML/CSS)
- Developed responsive kanban-style task dashboard with interactive status updates.
- Wrote unit tests using Django test runner with 85% code coverage.

EDUCATION
B.S. in Computer Science (2024)
Relevant Coursework: Data Structures, Algorithms, Database Systems, Web Development
"""
        },

        "data_analyst_ml_aspirant": {
            "title": "Data Analyst aspiring for AI/ML",
            "desc": "Experienced in Python, Pandas, NumPy, SQL, Tableau, with goal to advance into Machine Learning & LLMs.",
            "text": """Sarah Chen
Email: sarah.chen.analytics@example.com | Portfolio: sarahchen.io
Location: New York, NY

PROFESSIONAL SUMMARY
Data Analyst with 3 years of experience in data wrangling, exploratory data analysis, and predictive modeling using Python, SQL, and Tableau. Eager to transition into Machine Learning Engineering and Generative AI systems.

SKILLS & PROFICIENCIES
- Core: Python, SQL, Pandas, NumPy, Scikit-Learn, Matplotlib, Seaborn
- BI & Storage: Tableau, Power BI, PostgreSQL, Snowflake
- Emerging Interests: Machine Learning, Deep Learning, PyTorch, Natural Language Processing, LLMs

EXPERIENCE
Senior Business Data Analyst | RetailTech Global (2023 - Present)
- Built automated ETL pipelines in Python pulling from PostgreSQL and Snowflake, saving 12 hours weekly.
- Performed statistical cohort analysis and user churn predictive modeling using Scikit-Learn.
- Designed executive dashboards in Tableau tracking $15M+ annual recurring revenue (ARR).

Data Analyst | DataCorp Solutions (2021 - 2023)
- Analyzed marketing attribution across multi-channel campaigns using advanced SQL window functions.
- Developed regression models in Python to forecast quarterly inventory demand.

EDUCATION
B.S. in Applied Statistics & Informatics (2021)
"""
        },

        "cloud_devops_aspirant": {
            "title": "IT Admin / DevOps Aspirant",
            "desc": "Strong foundation in Linux, Bash scripting, Docker, AWS, networking, aiming for Cloud Architecture.",
            "text": """David Patel
Email: david.patel.cloud@example.com | GitHub: github.com/davidp-cloud
Location: Austin, TX

SUMMARY
Systems Administrator with 4 years of experience managing Linux servers, containerizing legacy apps with Docker, and provisioning cloud services in AWS. Seeking to advance to Cloud Solutions & DevOps Architect role with expertise in Kubernetes and Terraform.

TECHNICAL SKILLS
- Operating Systems & Scripting: Linux (Ubuntu, RHEL), Bash, Shell Scripting, Python
- Cloud & Virtualization: AWS (EC2, S3, IAM, VPC), Docker, Docker Compose
- Tools & Practices: Git, GitHub, CI/CD, Nginx, Prometheus, Network Security, Wireshark

EXPERIENCE
Systems Administrator | CloudSphere Networks (2022 - Present)
- Administered 50+ Linux production servers with 99.9% uptime SLA.
- Containerized 8 monolithic Python and PHP applications using Docker and multi-stage builds.
- Configured automated daily backup snapshots to AWS S3 and hardened IAM security policies.
- Implemented Prometheus server monitoring and custom metric alerts.

IT Infrastructure Specialist | Alpha Systems (2020 - 2022)
- Managed corporate network infrastructure, VLANs, VPN tunnels, and firewall rule configurations.
- Automated system provisioning routines using Bash and Python scripts.

CERTIFICATIONS
- AWS Certified Cloud Practitioner
- CompTIA Network+
"""
        }
    }
